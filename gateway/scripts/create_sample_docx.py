#!/usr/bin/env python3
"""
Generate sample.docx with PII (name, phone, email) for Gateway smoke tests.
Saves sample.docx, sample.docx.b64.txt, and prints base64 (single line) to stdout.
"""
from __future__ import annotations

import base64
import sys
from pathlib import Path

import docx


def create_sample_docx(output_path: Path) -> None:
    doc = docx.Document()
    doc.add_paragraph("Заявку подал Иванов Иван Иванович")
    doc.add_paragraph("")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Телефон"
    table.rows[0].cells[1].text = "Email"
    table.rows[1].cells[0].text = "+7 921 555-12-34"
    table.rows[1].cells[1].text = "ivanov@example.com"

    doc.add_paragraph("Документ подписан уполномоченным лицом.")
    doc.save(str(output_path))


def main() -> None:
    script_dir = Path(__file__).resolve().parent
    output_path = script_dir / "sample.docx"
    b64_path = script_dir / "sample.docx.b64.txt"

    create_sample_docx(output_path)
    print(f"Created: {output_path}", file=sys.stderr)

    encoded = base64.b64encode(output_path.read_bytes()).decode("ascii")
    b64_path.write_text(encoded, encoding="ascii")
    print(f"Saved base64: {b64_path}", file=sys.stderr)
    print(encoded)


if __name__ == "__main__":
    main()
